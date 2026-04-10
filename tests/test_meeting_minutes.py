"""End-to-end test for meeting minutes template."""

import os
import tempfile
import pytest
from pathlib import Path

from src.template.template import (
    load_config,
    parse_md_for_template,
    extract_sections,
    resolve_section_data,
)
from src.template.table_filler import fill_table_template
from src.core.parser import parse_md_to_elements
from src.core.renderer import render_to_docx
from src.core.extractor import extract_docx_to_elements
from src.core.md_writer import elements_to_md
from src.util.file_util import read, write


class TestMeetingMinutesE2E:
    """End-to-end test for meeting minutes template"""

    @pytest.fixture
    def template_path(self):
        """Path to the meeting minutes template"""
        return Path(__file__).parent.parent / "templates" / "meeting_minutes" / "template.docx"

    @pytest.fixture
    def config_path(self):
        """Path to the meeting minutes config"""
        return Path(__file__).parent.parent / "templates" / "meeting_minutes" / "config.yaml"

    @pytest.fixture
    def input_md_path(self):
        """Path to the test input markdown"""
        return Path(__file__).parent / "fixtures" / "meeting_input.md"

    @pytest.fixture
    def config(self, config_path):
        """Load the config"""
        return load_config(str(config_path))

    @pytest.fixture
    def md_content(self, input_md_path):
        """Read the markdown content"""
        return read(str(input_md_path))

    @pytest.fixture
    def parsed_data(self, config, md_content):
        """Parse markdown and resolve data"""
        front_matter, body = parse_md_for_template(md_content)
        sections = extract_sections(body)
        return resolve_section_data(config, front_matter, sections)

    def test_template_and_config_exist(self, template_path, config_path):
        """Test that template and config files exist"""
        assert template_path.exists(), f"Template not found: {template_path}"
        assert config_path.exists(), f"Config not found: {config_path}"

    def test_config_structure(self, config):
        """Test config has required structure"""
        assert "mode" in config
        assert config["mode"] == "table_fill"
        assert "fields" in config
        assert isinstance(config["fields"], list)
        assert "content_mapping" in config
        assert "headings" in config["content_mapping"]

    def test_markdown_parsing(self, md_content):
        """Test markdown is parsed correctly"""
        front_matter, body = parse_md_for_template(md_content)

        # Check front matter
        assert "meeting_time" in front_matter
        assert "meeting_location" in front_matter
        assert "recorder" in front_matter
        assert front_matter["meeting_time"] == "2026.4.10 14:00-16:00"

        # Check body has sections
        assert "与会人员" in body
        assert "会议摘要" in body
        assert "待明确事项" in body
        assert "下一步工作安排" in body

    def test_section_extraction(self, md_content):
        """Test sections are extracted correctly"""
        _, body = parse_md_for_template(md_content)
        sections = extract_sections(body)

        assert "与会人员" in sections
        assert "会议摘要" in sections
        assert "待明确事项" in sections
        assert "下一步工作安排" in sections

        # Check section content
        assert "地铁物业公司" in sections["与会人员"]
        assert "OA办公系统和公文功能" in sections["会议摘要"]
        assert "移动端APP" in sections["待明确事项"]
        assert "按以上意见调整" in sections["下一步工作安排"]

    def test_data_resolution(self, parsed_data):
        """Test data is resolved correctly from front matter and sections"""
        # Check front matter fields
        assert "meeting_time" in parsed_data
        assert "meeting_location" in parsed_data
        assert "recorder" in parsed_data
        assert parsed_data["meeting_time"] == "2026.4.10 14:00-16:00"
        assert parsed_data["meeting_location"] == "地铁物业305会议室"
        assert parsed_data["recorder"] == "张三、李四"

        # Check section fields
        assert "attendees" in parsed_data
        assert "summary_content" in parsed_data
        assert "pending_content" in parsed_data
        assert "next_steps_content" in parsed_data

        # Check section content
        assert "地铁物业公司" in parsed_data["attendees"]
        assert "OA办公系统和公文功能" in parsed_data["summary_content"]
        assert "移动端APP" in parsed_data["pending_content"]
        assert "按以上意见调整" in parsed_data["next_steps_content"]

    def test_fill_template(self, template_path, config, parsed_data):
        """Test filling the template with data"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            output_path = tmp.name

        try:
            # Fill template
            result = fill_table_template(str(template_path), config, parsed_data, output_path)

            # Verify output file was created
            assert os.path.exists(output_path)
            assert os.path.getsize(output_path) > 0
            assert result == output_path

            # Verify the filled document
            from docx import Document
            doc = Document(output_path)

            # Check table exists
            assert len(doc.tables) > 0
            table = doc.tables[0]

            # Check filled cells (by row, col)
            # Row 0: meeting_time (col 1)
            cell_text = table.rows[0].cells[1].text
            assert "2026.4.10" in cell_text, f"Expected time in cell [0,1], got: {cell_text}"

            # Row 0: meeting_location (col 3)
            cell_text = table.rows[0].cells[3].text
            assert "地铁物业305会议室" in cell_text or "305会议室" in cell_text, \
                f"Expected location in cell [0,3], got: {cell_text}"

            # Row 0: recorder (col 5)
            cell_text = table.rows[0].cells[5].text
            assert "张三" in cell_text, f"Expected recorder in cell [0,5], got: {cell_text}"

            # Row 1: attendees (col 1, spans 5 cols)
            cell_text = table.rows[1].cells[1].text
            assert "地铁物业公司" in cell_text, f"Expected attendees in cell [1,1], got: {cell_text}"

        finally:
            if os.path.exists(output_path):
                os.remove(output_path)

    def test_convert_flow(self, input_md_path):
        """Test MD -> DocElement -> Word conversion flow"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            output_path = tmp.name

        try:
            # Read markdown
            md_content = read(str(input_md_path))

            # Parse to elements
            elements = parse_md_to_elements(md_content)
            assert len(elements) > 0, "No elements parsed from markdown"

            # Render to docx
            render_to_docx(elements, output_path)

            # Verify output
            assert os.path.exists(output_path)
            assert os.path.getsize(output_path) > 0

            # Verify content
            from docx import Document
            doc = Document(output_path)
            text = "\n".join(p.text for p in doc.paragraphs)

            assert "会议议程" in text or "OA系统" in text, \
                f"Expected content in converted docx, got: {text[:200]}"

        finally:
            if os.path.exists(output_path):
                os.remove(output_path)

    def test_extract_flow(self, input_md_path):
        """Test MD -> DocElement -> Word -> DocElement -> MD roundtrip"""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Step 1: Convert MD to DOCX
            docx_path = os.path.join(tmpdir, "intermediate.docx")
            md_content = read(str(input_md_path))
            elements = parse_md_to_elements(md_content)
            render_to_docx(elements, docx_path)

            # Step 2: Extract back to MD
            extracted_md_path = os.path.join(tmpdir, "extracted.md")
            extracted_elements, images = extract_docx_to_elements(docx_path)
            extracted_md = elements_to_md(extracted_elements)

            # Write and verify
            write(extracted_md_path, extracted_md)

            # Verify extracted content
            assert os.path.exists(extracted_md_path)
            assert os.path.getsize(extracted_md_path) > 0

            # Check key content is preserved
            extracted_content = read(extracted_md_path)
            assert "会议" in extracted_content or "OA" in extracted_content, \
                f"Key content not preserved in extraction: {extracted_content[:200]}"

    def test_full_roundtrip(self, template_path, config_path, input_md_path):
        """Test complete roundtrip: MD -> Template Fill -> Extract -> MD"""
        with tempfile.TemporaryDirectory() as tmpdir:
            # Step 1: Fill template
            filled_path = os.path.join(tmpdir, "filled.docx")
            config = load_config(str(config_path))
            md_content = read(str(input_md_path))

            front_matter, body = parse_md_for_template(md_content)
            sections = extract_sections(body)
            data = resolve_section_data(config, front_matter, sections)

            fill_table_template(str(template_path), config, data, filled_path)

            # Verify filled docx
            assert os.path.exists(filled_path)
            assert os.path.getsize(filled_path) > 0

            # Step 2: Extract from filled docx
            extracted_elements, images = extract_docx_to_elements(filled_path)
            extracted_md = elements_to_md(extracted_elements)

            # Verify extraction
            assert len(extracted_md) > 0

            # Check that key data is in the extracted content
            assert "2026.4.10" in extracted_md or "14:00-16:00" in extracted_md, \
                "Meeting time not found in extracted content"
            assert "张三" in extracted_md or "李四" in extracted_md, \
                "Recorder not found in extracted content"

    def test_cli_template_list(self):
        """Test CLI template list command"""
        import subprocess
        result = subprocess.run(
            ["poetry", "run", "python", "main.py", "template", "list"],
            capture_output=True,
            text=True,
            cwd=Path(__file__).parent.parent
        )
        assert result.returncode == 0, f"CLI failed: {result.stderr}"
        assert "meeting_minutes" in result.stdout, \
            f"meeting_minutes template not found in list: {result.stdout}"

    def test_cli_fill_command(self, template_path, config_path, input_md_path):
        """Test CLI fill command"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            output_path = tmp.name

        try:
            import subprocess
            result = subprocess.run(
                [
                    "poetry", "run", "python", "main.py", "fill",
                    str(input_md_path),
                    "-t", str(template_path),
                    "-c", str(config_path),
                    "-o", output_path
                ],
                capture_output=True,
                text=True,
                cwd=Path(__file__).parent.parent
            )

            assert result.returncode == 0, f"CLI fill failed: {result.stderr}"
            assert os.path.exists(output_path), "Output file not created"
            assert os.path.getsize(output_path) > 0, "Output file is empty"

        finally:
            if os.path.exists(output_path):
                os.remove(output_path)

    def test_cli_convert_command(self, input_md_path):
        """Test CLI convert command"""
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as tmp:
            output_path = tmp.name

        try:
            import subprocess
            result = subprocess.run(
                [
                    "poetry", "run", "python", "main.py", "convert",
                    str(input_md_path),
                    "-o", output_path
                ],
                capture_output=True,
                text=True,
                cwd=Path(__file__).parent.parent
            )

            assert result.returncode == 0, f"CLI convert failed: {result.stderr}"
            assert os.path.exists(output_path), "Output file not created"
            assert os.path.getsize(output_path) > 0, "Output file is empty"

        finally:
            if os.path.exists(output_path):
                os.remove(output_path)
