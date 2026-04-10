import shutil
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def fill_table_template(
    template_path: str,
    config: dict,
    data: dict[str, str],
    output_path: str,
) -> str:
    shutil.copy(template_path, output_path)
    doc = Document(output_path)

    for field in config["fields"]:
        name = field["name"]
        if name not in data:
            continue

        table_idx = field.get("table", 0)
        row_idx = field["row"]
        col_idx = field["col"]

        table = doc.tables[table_idx]
        cell = table.rows[row_idx].cells[col_idx]
        _set_cell_text_preserve_format(cell, data[name])

    doc.save(output_path)
    return output_path


def _set_cell_text_preserve_format(cell, new_text: str) -> None:
    paragraphs = cell.paragraphs
    if not paragraphs:
        return

    for p in paragraphs:
        for run in list(p.runs):
            run._r.getparent().remove(run._r)
        for child in list(p._p):
            if child.tag != qn("w:pPr"):
                p._p.remove(child)

    lines = new_text.split("\n")
    if lines:
        run = paragraphs[0].add_run(lines[0])
        _set_run_font(run)
        for line in lines[1:]:
            new_p = OxmlElement("w:p")
            new_r = OxmlElement("w:r")
            new_t = OxmlElement("w:t")
            new_t.text = line
            new_t.set(qn("xml:space"), "preserve")
            new_r.append(new_t)
            new_p.append(new_r)
            cell._tc.append(new_p)


def _set_run_font(run) -> None:
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "宋体")
    rPr.append(rFonts)
