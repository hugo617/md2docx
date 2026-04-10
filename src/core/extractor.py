import os
import zipfile
from docx import Document
from docx.oxml.ns import qn

from .elements import DocElement, ElementType


def extract_docx_to_elements(docx_path: str) -> tuple[list[DocElement], dict[str, str]]:
    doc = Document(docx_path)
    img_dir = os.path.splitext(docx_path)[0] + "_images"
    images = _extract_images(docx_path, img_dir)

    elements = []
    for element in doc.element.body:
        tag = element.tag.split("}")[-1]
        if tag == "p":
            para = _find_paragraph(doc, element)
            if para:
                el = _extract_paragraph(para, images)
                if el:
                    elements.append(el)
        elif tag == "tbl":
            table = _find_table(doc, element)
            if table:
                elements.append(_extract_table(table))

    return elements, images


def _extract_images(docx_path: str, output_dir: str) -> dict[str, str]:
    images = {}
    os.makedirs(output_dir, exist_ok=True)
    with zipfile.ZipFile(docx_path, "r") as z:
        for name in z.namelist():
            if name.startswith("word/media/"):
                data = z.read(name)
                filename = os.path.basename(name)
                with open(os.path.join(output_dir, filename), "wb") as f:
                    f.write(data)
                images[name] = os.path.join(output_dir, filename)
    return images


def _extract_paragraph(para, images: dict) -> DocElement | None:
    style_name = para.style.name if para.style else "Normal"
    text = para.text.strip()
    if not text:
        return None
    if style_name.startswith("Heading"):
        level = int(style_name.split()[-1]) if style_name[-1].isdigit() else 1
        return DocElement(element_type=ElementType.HEADING, level=level, text=text)
    if "List" in style_name:
        list_type = ElementType.ORDERED_LIST if "Number" in style_name else ElementType.UNORDERED_LIST
        formatted = _get_run_formatting(para)
        return DocElement(element_type=ElementType.LIST_ITEM, text=formatted, style_name=list_type.value)
    if _is_code_paragraph(para):
        return DocElement(element_type=ElementType.CODE_BLOCK, text=text)
    if "Quote" in style_name:
        return DocElement(element_type=ElementType.BLOCKQUOTE, text=text)
    children = _extract_inline_runs(para)
    return DocElement(element_type=ElementType.PARAGRAPH, children=children)


def _extract_table(table) -> DocElement:
    rows = []
    for row in table.rows:
        seen = set()
        cells = []
        for cell in row.cells:
            tc_id = id(cell._tc)
            if tc_id in seen:
                continue
            seen.add(tc_id)
            cells.append(DocElement(element_type=ElementType.TABLE_CELL, text=cell.text.strip()))
        rows.append(DocElement(element_type=ElementType.TABLE_ROW, children=cells))
    return DocElement(element_type=ElementType.TABLE, children=rows)


def _get_run_formatting(para) -> str:
    parts = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold and run.italic:
            parts.append(f"***{text}***")
        elif run.bold:
            parts.append(f"**{text}**")
        elif run.italic:
            parts.append(f"*{text}*")
        else:
            parts.append(text)
    return "".join(parts)


def _is_code_paragraph(para) -> bool:
    for run in para.runs:
        if run.font.name and run.font.name.lower() in ("consolas", "courier new"):
            return True
    pPr = para._p.find(qn("w:pPr"))
    if pPr is not None:
        shd = pPr.find(qn("w:shd"))
        if shd is not None and shd.get(qn("w:fill")):
            return True
    return False


def _extract_inline_runs(para) -> list[DocElement]:
    elements = []
    for run in para.runs:
        text = run.text
        if not text:
            continue
        if run.bold:
            elements.append(DocElement(element_type=ElementType.BOLD, text=text))
        elif run.italic:
            elements.append(DocElement(element_type=ElementType.ITALIC, text=text))
        else:
            elements.append(DocElement(element_type=ElementType.PARAGRAPH, text=text))
    return elements if elements else [DocElement(element_type=ElementType.PARAGRAPH, text=para.text)]


def _find_paragraph(doc, element):
    for p in doc.paragraphs:
        if p._p is element:
            return p
    return None


def _find_table(doc, element):
    for t in doc.tables:
        if t._tbl is element:
            return t
    return None
