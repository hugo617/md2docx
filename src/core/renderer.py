from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .elements import DocElement, ElementType
from .style_map import DEFAULT_STYLE_CONFIG


def render_to_docx(
    elements: list[DocElement],
    output_path: str,
    style_config: dict | None = None,
) -> str:
    config = style_config or DEFAULT_STYLE_CONFIG
    doc = Document()
    _apply_default_styles(doc, config)

    for element in elements:
        _render_element(doc, element, config)

    doc.save(output_path)
    return output_path


def _apply_default_styles(doc: Document, config: dict) -> None:
    style = doc.styles["Normal"]
    style.font.name = config["fonts"]["body"]
    style.font.size = Pt(config["sizes"]["body"])
    rPr = style.element.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), config["fonts"]["body_east_asia"])
    rPr.append(rFonts)

    for level in range(1, 4):
        hstyle = doc.styles[f"Heading {level}"]
        hstyle.font.name = config["fonts"]["heading"]
        hstyle.font.bold = True
        hrPr = hstyle.element.get_or_add_rPr()
        hrFonts = OxmlElement("w:rFonts")
        hrFonts.set(qn("w:eastAsia"), config["fonts"]["heading_east_asia"])
        hrPr.append(hrFonts)


def _render_element(doc: Document, element: DocElement, config: dict) -> None:
    match element.element_type:
        case ElementType.HEADING:
            doc.add_heading(element.text, level=element.level)

        case ElementType.PARAGRAPH:
            p = doc.add_paragraph()
            if element.children:
                for child in element.children:
                    _render_inline(p, child)
            elif element.text:
                p.add_run(element.text)

        case ElementType.CODE_BLOCK:
            _render_code_block(doc, element, config)

        case ElementType.UNORDERED_LIST | ElementType.ORDERED_LIST:
            style = "List Bullet" if element.element_type == ElementType.UNORDERED_LIST else "List Number"
            for item in element.children:
                doc.add_paragraph(item.text, style=style)

        case ElementType.BLOCKQUOTE:
            _render_blockquote(doc, element, config)

        case ElementType.TABLE:
            _render_table(doc, element, config)

        case ElementType.HORIZONTAL_RULE:
            _render_hr(doc)

        case ElementType.IMAGE:
            p = doc.add_paragraph()
            run = p.add_run(f"[图片: {element.url}]")
            run.font.color.rgb = RGBColor(128, 128, 128)


def _render_inline(p, element: DocElement) -> None:
    match element.element_type:
        case ElementType.BOLD:
            run = p.add_run(element.text)
            run.bold = True
        case ElementType.ITALIC:
            run = p.add_run(element.text)
            run.italic = True
        case ElementType.STRIKETHROUGH:
            run = p.add_run(element.text)
            run.font.strike = True
        case ElementType.CODE_INLINE:
            run = p.add_run(element.text)
            run.font.name = "Consolas"
            run.font.size = Pt(9)
        case ElementType.LINK:
            run = p.add_run(element.text)
            run.font.color.rgb = RGBColor(0, 0, 238)
            run.underline = True
        case _:
            if element.text:
                p.add_run(element.text)


def _render_code_block(doc, element: DocElement, config: dict) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), config["colors"]["code_bg"])
    shading.set(qn("w:val"), "clear")
    pPr.append(shading)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "360")
    ind.set(qn("w:right"), "360")
    pPr.append(ind)
    run = p.add_run(element.text)
    run.font.name = "Consolas"
    run.font.size = Pt(config["sizes"]["code"])


def _render_blockquote(doc, element: DocElement, config: dict) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    pPr.append(ind)
    pBdr = OxmlElement("w:pBdr")
    left_bdr = OxmlElement("w:left")
    left_bdr.set(qn("w:val"), "single")
    left_bdr.set(qn("w:sz"), "12")
    left_bdr.set(qn("w:space"), "4")
    left_bdr.set(qn("w:color"), config["colors"]["quote_border"])
    pBdr.append(left_bdr)
    pPr.append(pBdr)
    p.add_run(element.text)


def _render_table(doc, element: DocElement, config: dict) -> None:
    rows_data = []
    for row_el in element.children:
        cells = [c.text for c in row_el.children]
        rows_data.append(cells)
    if not rows_data:
        return
    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.style = "Table Grid"
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                if i == 0:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            run.bold = True
                    _set_cell_shading(cell, config["colors"]["table_header_bg"])


def _render_hr(doc) -> None:
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "CCCCCC")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_cell_shading(cell, color: str) -> None:
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    shading.set(qn("w:val"), "clear")
    cell._tc.get_or_add_tcPr().append(shading)
